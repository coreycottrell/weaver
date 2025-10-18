#!/usr/bin/env python3
"""
Validation Test: doc-synthesizer PDF + DOCX skills
Tests ability to extract from PDF and create DOCX
"""

import pdfplumber
from docx import Document
from pathlib import Path

print("🧬 doc-synthesizer: Validating PDF + DOCX skills")
print("=" * 60)

# Test 1: PDF Extraction
print("\n📄 TEST 1: PDF Extraction")
try:
    with pdfplumber.open("test_skills_invoice.pdf") as pdf:
        first_page = pdf.pages[0]
        text = first_page.extract_text()
        tables = first_page.extract_tables()

        print(f"✅ PDF opened successfully")
        print(f"✅ Text extracted: {len(text)} characters")
        print(f"✅ Tables found: {len(tables)}")
        print(f"\nSample text (first 200 chars):")
        print(text[:200] if text else "No text found")

        pdf_success = True
except Exception as e:
    print(f"❌ PDF extraction failed: {e}")
    pdf_success = False

# Test 2: DOCX Creation
print("\n📝 TEST 2: DOCX Creation")
try:
    doc = Document()
    doc.add_heading('Skills Validation Test', 0)
    doc.add_paragraph('This document was created by doc-synthesizer using the DOCX skill.')

    # Add some content from the PDF
    if pdf_success and text:
        doc.add_heading('Extracted from PDF:', level=1)
        doc.add_paragraph(text[:500])  # First 500 chars

    output_path = "test_doc_synthesizer_output.docx"
    doc.save(output_path)

    # Verify file was created
    if Path(output_path).exists():
        file_size = Path(output_path).stat().st_size
        print(f"✅ DOCX created successfully")
        print(f"✅ Output file: {output_path}")
        print(f"✅ File size: {file_size} bytes")
        docx_success = True
    else:
        print(f"❌ DOCX file not found after save")
        docx_success = False

except Exception as e:
    print(f"❌ DOCX creation failed: {e}")
    docx_success = False

# Final Results
print("\n" + "=" * 60)
print("📊 VALIDATION RESULTS:")
print(f"  PDF skill:  {'✅ PASS' if pdf_success else '❌ FAIL'}")
print(f"  DOCX skill: {'✅ PASS' if docx_success else '❌ FAIL'}")
print(f"\n  Overall:    {'✅ READY FOR PRODUCTION' if (pdf_success and docx_success) else '❌ NOT READY'}")
print("=" * 60)

exit(0 if (pdf_success and docx_success) else 1)
